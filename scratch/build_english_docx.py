# -*- coding: utf-8 -*-
import sys
import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def build_docx(txt_path, docx_path):
    print(f"Reading {txt_path} and compiling Word document...")
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: {txt_path} not found.")
        return False

    doc = Document()

    # Set document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11.5)

    # Split content by page markers: [PAGE N - ...]
    page_blocks = re.split(r'(\[PAGE \d+ - [^\]]+\])', content)
    
    if len(page_blocks) < 3:
        p = doc.add_paragraph(content)
        doc.save(docx_path)
        print(f"Compiled without page markers. Saved to {docx_path}")
        return True

    header_text = page_blocks[0].strip()
    if header_text:
        doc.add_paragraph(header_text)
        doc.add_page_break()

    # Keep track of whether the last line added was Devanagari
    last_was_dev = False

    for i in range(1, len(page_blocks), 2):
        page_header = page_blocks[i].strip()
        page_body = page_blocks[i+1] if (i+1) < len(page_blocks) else ""
        page_body = page_body.strip()

        # Add Page Heading
        h = doc.add_paragraph()
        run = h.add_run(page_header)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Georgia'
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

        # Add Body paragraphs
        paragraphs = page_body.split('\n')
        
        last_was_dev = False
        
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            
            # Check line type:
            is_dev = bool(re.search(r'[\u0900-\u097F]', p_text))
            
            if p_text.startswith("Source:") or p_text.startswith("Source "):
                # Footer source page
                run = p.add_run(p_text)
                run.italic = True
                run.font.size = Pt(9.5)
                run.font.name = 'Georgia'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                last_was_dev = False
            elif is_dev:
                # Devanagari Sanskrit Verse
                run = p.add_run(p_text)
                run.bold = True
                run.font.name = 'Nirmala UI'  # Excellent Devanagari font
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                last_was_dev = True
            elif last_was_dev:
                # IAST Transliteration (follows Devanagari)
                run = p.add_run(p_text)
                run.italic = True
                run.font.name = 'Georgia'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(158, 58, 26)  # Rust color #9e3a1a
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                last_was_dev = False
            elif p_text.startswith("[") and p_text.endswith("]"):
                # Verse translation block
                run = p.add_run(p_text)
                run.italic = True
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                last_was_dev = False
            elif p_text.startswith("✦"):
                # Separator
                run = p.add_run(p_text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_was_dev = False
            else:
                # Regular English commentary text
                p.add_run(p_text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                last_was_dev = False

        # Add page break if it's not the last page
        if i + 2 < len(page_blocks):
            doc.add_page_break()

    doc.save(docx_path)
    print(f"Compilation complete! Saved to {docx_path}")
    return True

if __name__ == "__main__":
    txt_path = "katopanishad-english.txt"
    docx_path = "katopanishad-english.docx"
    build_docx(txt_path, docx_path)
