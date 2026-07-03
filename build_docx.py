# -*- coding: utf-8 -*-
import sys
import re
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def build_docx(txt_path, docx_path):
    print(f"Reading {txt_path} and compiling Word document...")
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: {txt_path} not found.")
        return False

    doc = Document()

    # Set document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base styling (Use a clean unicode font supporting Kannada)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Nirmala UI'
    font.size = Pt(12)

    # Split content by page markers: [PAGE N - ...]
    page_blocks = re.split(r'(\[PAGE \d+ - [^\]]+\])', content)
    
    if len(page_blocks) < 3:
        # If no page markers, just dump everything
        p = doc.add_paragraph(content)
        doc.save(docx_path)
        print(f"Compiled without page markers. Saved to {docx_path}")
        return True

    # First block is metadata or empty text before the first page marker
    header_text = page_blocks[0].strip()
    if header_text:
        doc.add_paragraph(header_text)
        doc.add_page_break()

    for i in range(1, len(page_blocks), 2):
        page_header = page_blocks[i].strip()
        page_body = page_blocks[i+1] if (i+1) < len(page_blocks) else ""
        page_body = page_body.strip()

        # Add Page Heading
        h = doc.add_paragraph()
        run = h.add_run(page_header)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Nirmala UI'
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

        # Add Body paragraphs
        paragraphs = page_body.split('\n')
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            
            # Highlight footer or sources
            if p_text.startswith("ಆಧಾರ:"):
                run = p.add_run(p_text)
                run.italic = True
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif p_text.startswith("||") or p_text.endswith("||") or p_text.startswith("✦"):
                # Centered verses or section dividers
                run = p.add_run(p_text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.add_run(p_text)

        # Add page break if it's not the last page
        if i + 2 < len(page_blocks):
            doc.add_page_break()

    doc.save(docx_path)
    print(f"Compilation complete! Saved to {docx_path}")
    return True

if __name__ == "__main__":
    txt_path = "kata-upanishad-text.txt"
    docx_path = "kata-upanishad.docx"
    build_docx(txt_path, docx_path)
